"""Read text from PDF and DOCX resume files."""

from pathlib import Path

import fitz  # PyMuPDF
from docx import Document


def read_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF using PyMuPDF."""
    text_parts: list[str] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts)


def read_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX file using python-docx."""
    from io import BytesIO

    document = Document(BytesIO(file_bytes))
    paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]

    # Also read text from tables (common in resume templates)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)

    return "\n".join(paragraphs)


def read_resume(file_bytes: bytes, filename: str) -> str:
    """
    Route to the correct reader based on file extension.

    Raises:
        ValueError: If the file type is not supported.
    """
    extension = Path(filename).suffix.lower()

    if extension == ".pdf":
        return read_pdf(file_bytes)
    if extension in {".docx", ".doc"}:
        if extension == ".doc":
            raise ValueError(
                "Legacy .doc files are not supported. Please save as .docx and re-upload."
            )
        return read_docx(file_bytes)

    raise ValueError(f"Unsupported file type: {extension}. Use PDF or DOCX.")
